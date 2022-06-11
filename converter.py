import os
import re
import docx
from docx import Document

# load examples from a text

def get_exes(corp):
        
    gla_rgx = '''^ ([^а-яА-Я].+?)$'''
    glb_rgx = '''^ ([а-яА-Я].+?)$'''
    glft_rgx = '''^([а-яА-Я].+?)$'''

    exes = []
    ex = dict.fromkeys(['gla', 'glb', 'glft', 'comm'], [])
    nex = 'gla'

    for line in corp.splitlines():
        # the text
        if bool(re.search(gla_rgx, line)) and (nex == 'gla'):
            s = ' '.join(re.findall(gla_rgx, line)[0].split('\t|'))
            ex['gla'] = ex['gla'] + [s]
            nex = 'glb'
        # the word-by-word glosses
        elif bool(re.search(glb_rgx, line)) and (nex == 'glb'):
            s = ' '.join(re.findall(glb_rgx, line)[0].split('\t|'))
            ex['glb'] = ex['glb'] + [s]
            nex = 'gla'
        # the translation
        elif bool(re.search(glft_rgx, line)) and (nex == 'gla'):
            ex['glft'] = ex['glft'] + re.findall(glft_rgx, line)
            nex = 'comm'
        # comments if there are such
        elif nex == 'comm':
            ex['comm'] = ex['comm'] + re.findall(glft_rgx, line)
            nex = 'gla'
            exes.append(ex)
            ex = dict.fromkeys(['gla', 'glb', 'glft', 'comm'], [])

    for ex in exes:
        for key in ex:
            ex[key] = ' '.join(ex[key])
            
    return exes

# get metadata of a text

def get_meta(text):
    
    params = ['title', 'title-abbreviation', 'source', 'comment']
    meta = dict.fromkeys(params, [])

    for par in meta:
        meta_rgx = f'''{par}\t(.*)'''
        meta[par] = meta[par] + re.findall(meta_rgx, text)
        
    return meta

# save the corpus output as a dictionary

def load_exes(filename):
    
    texts = []
    res = []
    
    with open(filename, 'r', encoding='utf8') as f:
        corp = f.read()
        
    for text in corp.split('kh_')[1:]:
        texts.append('title\tkh_' + text)
            
    for text in texts:
        text_dic = dict()
        text_dic['meta'] = get_meta(text)
        text_dic['exes'] = get_exes(text)
        res.append(text_dic)
        
    return res

# write the dict to a docx file

def print_docx(in_fn, out_fn, style_fn="template.docx"):
    
    cur_dir = [f for f in os.listdir('.') if os.path.isfile(f)]
    
    if style_fn not in cur_dir:
        style_fn = input('Style template file not found. Please insert the path to the style template .docx file: ')
    
    doc = docx.Document(style_fn)

    for text in load_exes(in_fn):
        doc.add_paragraph(text['meta']['title'][1], style='Heading 1')
        if text['meta']['source'] != []:
            doc.add_paragraph('\n'.join(text['meta']['source']), style='Normal')
        if text['meta']['comment'] != []:
            doc.add_paragraph('\n'.join(text['meta']['comment']), style='Normal')
        for ex in text['exes']:
            doc.add_paragraph(ex['gla'], style='ex_gla')
            doc.add_paragraph(ex['glb'], style='ex_glb')
            glft = ex['glft']
            doc.add_paragraph(f'`{glft}\'', style='ex_glft')
            if ex['comm']:
                doc.add_paragraph(ex['comm'], style='ex_comment')

    doc.save(out_fn)
    
if __name__ == "__main__":
    print('''
        This is a FLEx-to-Word converter.
        ''')
    in_fn = input('Insert path to plain text file containing the FLEx export: ')
    out_fn = input('Insert output filename: ')
    print_docx(in_fn, out_fn, style_fn="template.docx")
    print('File ready!')
